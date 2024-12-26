import openai
import pandas as pd
from docx import Document
from openai import AzureOpenAI
from io import BytesIO
from urllib.parse import urlparse
import mimetypes
from mimetypes import guess_type, add_type
import base64
import re
import os 
import time
import streamlit as st
import zipfile




class DocumentGeneration:
    def __init__(self, template_path, data, template):

        self.client = AzureOpenAI(
        azure_endpoint=st.secrets["AZURE_ENDPOINT"],
        api_version=st.secrets["AZURE_VERSION"],
        api_key=st.secrets["AZURE_KEY"])

        self.model_name = st.secrets["AZURE_MODEL"]
        self.prompt = open("prompt.txt", "r").read()
        self.template_path = template_path
        self.template = template
        self.data = data
    
    def get_prompt(self, context, input_data):
        return f"""
        You are an AI tasked with replacing placeholders in an HR document template using the provided input data.  

        Template : {context}  
        Input Data : {input_data.to_csv(index=False)} 
        
        """ + self.prompt
    
    def call_gpt(self, prompt):

        try:
            completion = self.client.chat.completions.create(
                model=self.model_name,
                temperature=0.1,
                messages=[{'role': 'system', 'content': 'You are a helpful assistant.'},
                          {"role": "user", "content": prompt}])

            output = completion.choices[0].message.content
            return output, True
        except Exception as e:
            print(f"Error calling GPT API: {e}")
            return e, False

    def gpt_fill_placeholder(self, context, input_data):

        prompt = self.get_prompt(context=context, input_data=input_data)
        output, gpt_flag = self.call_gpt(prompt)
        if gpt_flag:
            return output


    def get_full_name_gpt(self, data):
        prompt = f"""
        Read the input CSV and provide the full name:
        - If a 'Full Name' column exists, return its value.
        - If 'First Name' and 'Last Name' columns exist, combine them as 'First Name Last Name'.
        - If only 'First Name' exists, return its value.

        Input CSV:
        {data.to_csv(index=False)}

        Do NOT add any introductory or concluding phrases.
        """

        output = self.call_gpt(prompt)
        return output

    def get_table_value(self, placeholder, data):
        prompt = f"""
        Read the input CSV and provide the relevant and appropriate value for {placeholder}:
        Answer just the value.
        If you don't find relevant and approprioate value, just return NA

        Input CSV:
        {data.to_csv(index=False)}

        Do NOT add any introductory or concluding phrases.
        """

        output, gpt_flag = self.call_gpt(prompt)
        if gpt_flag:
            return output

    @staticmethod
    def check_for_placeholders(text):
        """Check if a string contains placeholders."""
        text = text.lower().replace(' /','/').replace('/ ','/')
        #"date:", "place:", "location:",
        return any(
            placeholder in text for placeholder in ["<<", ">>", "{", "}", "__",  "his/her", "her/his", "him/her", "her/him", "mr/mrs", "mr/ms"])

    @staticmethod
    def get_signature_full_name(text):
        pattern = r"<<.*?>>"
        # Find all matches
        matches = re.findall(pattern, text)
        if matches:
            return matches[0]


    def doc_gen(self, data, file_name):
        doc = Document(self.template_path)
        full_name = file_name
        i = 1

        for para in doc.paragraphs:
            print("para:   ", i)
            print("papa text:  ", para.text)
            placeholder_present = self.check_for_placeholders(para.text)

            if placeholder_present:
                if '    <<' not in para.text.lower():
                    gpt_output = self.gpt_fill_placeholder(para.text, data)
                    para.text = gpt_output
                elif '    <<' in para.text.lower() and 'name' in para.text.lower():
                    full_name_placeholder = self.get_signature_full_name(para.text)
                    para.text ='##' + para.text.replace(full_name_placeholder, full_name) + '##'

            i += 1

        for table in doc.tables:
            # Ensure the table has exactly two columns
            if all(len(row.cells) == 2 for row in table.rows):
                for row in table.rows:
                    key_cell, value_cell = row.cells
                    print(key_cell.text, value_cell.text)
                    # Check if the second cell (value_cell) is empty
                    if not value_cell.text.strip():  # Second cell is empty
                        placeholder = key_cell.text.strip()  # Get the text from the first cell
                        # Pass the placeholder and context to GPT to generate the value
                        updated_value = self.get_table_value(placeholder, data)
                        # Update the second cell with GPT's output
                        value_cell.text = updated_value


        for paragraph in doc.paragraphs:
            # Iterate over a copy of the runs to avoid modifying the list while iterating
            original_runs = list(paragraph.runs)

            for run in original_runs:
                # Skip runs that contain images or other non-text content
                if run.text.strip() == "":
                    continue  # Likely contains non-text content, such as an image or hyperlink

                # Split the run text into segments based on the ## markers
                print("run.text:   ", run)
                matches = re.split(r'(##.*?##)', run.text)
                print("matches:  ", matches)
                # Clear the original text in the run
                run.text = ""

                # Process matches and recreate the runs
                for match in matches:
                    if match.startswith("##") and match.endswith("##"):
                        print("if")
                        bold_text = match.strip("#")
                        new_run = paragraph.add_run(bold_text)
                        new_run.bold = True
                    elif match.startswith('"') and match.endswith('"'):
                        print("elif")
                        bold_text = match
                        new_run = paragraph.add_run(bold_text)
                        new_run.bold = True
                    else:
                        print("else")
                        # Retain the original formatting for non-bold text
                        new_run = paragraph.add_run(match)
                        new_run.font.name = run.font.name
                        new_run.font.size = run.font.size
                        new_run.font.color.rgb = run.font.color.rgb
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.bold = run.bold  # Retain existing bold formatting if present

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer, file_name

    
    def doc_gen_main(self):

        documents = []

        no_of_rows = self.data.shape[0]
        for i in range(no_of_rows):
            df_new = self.data.iloc[[i],:]
            full_name, gpt_flag = self.get_full_name_gpt(df_new)
            file_buffer, filename = self.doc_gen(df_new, full_name)
            documents.append((file_buffer, filename))
            st.write(f"{i+1}:  Document generation completed for {filename}")

        # Step 2: Handle single or multiple files
        if len(documents) == 1:
            # Single file - Prepare for direct download
            file_buffer, filename = documents[0]
            return file_buffer, self.template + '_' + filename + '.docx'
        else:
            # Multiple files - Bundle into a ZIP archive
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                for file_buffer, filename in documents:
                    # Ensure the internal files in the ZIP have a .docx extension
                    if not filename.endswith(".docx"):
                        filename += ".docx"
                        filename = self.template + '_' + filename
                    zip_file.writestr(filename, file_buffer.read())
            zip_buffer.seek(0)
            return zip_buffer, self.template + '_''documents.zip'
            


# if __name__ == "__main__":
#     template = 'templates/Offshore_Appointment letter_with Relocation.docx'
#     data = '/home/affine/Projects/HR-Tools/test_ouput/Book 4.xlsx'
#     obj = DocumentGeneration(template_path=template, data=data)
#     a, b = obj.doc_gen_main()
#     print(a, b)
