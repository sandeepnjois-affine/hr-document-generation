import openai
import pandas as pd
from docx import Document
from openai import AzureOpenAI
from io import BytesIO
from urllib.parse import urlparse
import mimetypes
from mimetypes import guess_type, add_type
import base64
import re
import os 
import time
import streamlit as st

client = AzureOpenAI(
    azure_endpoint=st.secrets["AZURE_ENDPOINT"],
    api_version=st.secrets["AZURE_VERSION"],
    api_key=st.secrets["AZURE_KEY"])

model_name = st.secrets["AZURE_MODEL"]



import re

def wrap_double_quotes_with_hashes(text):
    # Regex to find phrases within double quotes “ ”
    matches = re.findall(r'“(.*?)”', text)

    # Replace each match with ##“phrase”##
    for match in matches:
        original = f'“{match}”'
        wrapped = f'##“{match}”##'
        text = text.replace(original, wrapped)

    return text

# # Example usage
# input_text = 'The “Company” has a policy that ensures “Employee Satisfaction” is a priority.'
# updated_text = wrap_double_quotes_with_hashes(input_text)

# print("Updated Text:")
# print(updated_text)


def wrap_phrases_with_hashes(text):
    # Regular expression to match phrases starting with a capital letter and ending with a colon,
    # but not already wrapped in double hashes
    # matches = re.findall(r'(?<!##)([A-Z][a-zA-Z\s]*:)(?!##)', text)
    matches = re.findall(r'(?<!##)([A-Z][a-zA-Z\s]*:)', text)
    # Replace each match by wrapping it within double hashes
    for match in matches:
        wrapped_phrase = f"##{match.strip()}##"
        text = text.replace(match, wrapped_phrase)
    return text

def populate_docx_with_gpt(template_path, data, gpt_model="gpt-4"):
    """
    Populate a .docx template with values from an input file using GPT-4.

    Args:
        template_path (str): Path to the .docx template file.
        input_data_path (str): Path to the input CSV or Excel file.
        output_path (str): Path to save the updated .docx file.
        gpt_model (str): GPT model to use for API calls (default: "gpt-4").
    """
    # Load the input data (supports both CSV and Excel)
    # if input_data_path.endswith(".csv"):
    #     data = pd.read_csv(input_data_path)
    # elif input_data_path.endswith((".xlsx", ".xls")):
    #     data = pd.read_excel(input_data_path)
    # else:
    #     raise ValueError("Input file must be a CSV or Excel file.")

    # Convert the input data to a dictionary of attributes and values
    # input_data = data.to_dict(orient="records")[0]  # Assuming single-row input

    # Load the template .docx file
    doc = Document(template_path)

    print('type data:' , type(data))
    print("data:  ", data)
    # Function to call GPT-4 for placeholder filling
    def gpt_fill_placeholder(context, input_data):
        #     prompt = f"""
        # You are an AI assisting with HR document generation.
        # Below is the context for an HR template and the input data attributes.
        # Context template with some placeholder(s): {context}
        # Input data attributes: {input_data.to_csv(index=False)}

        # Replace the placeholder with the most appropriate value by understanding the context around the placeholder.

        # Do NOT add any introductory or concluding phrases. Just answer what is needed.
        # If you think there is no relevant information to be replaced with any placeholder, keep the placeholder as it is.

        # """
        # prompt = f"""
        # You are an AI tasked with replacing placeholders in an HR document template based on the provided input data. 
        # Template: {context}
        # Input Data: {input_data.to_csv(index=False)}

        # Replace each placeholder with the most contextually relevant value from the input data. 
        # - If no relevant value exists, leave the placeholder unchanged. 
        # - Provide only the updated text or table content, without any additional commentary.
        # - Based on the gender of the name provided in the input data format the pronoun (his/her, mr/mrs) accordingly.
        # - If there is date placeholder like <<LWD dd/MM/YYYY>> then modify the date as 12/12/2024 (This is an example).

        # If the place holders are within the << >> or in the form of a __________. Replace it with the value
        # If the place holder is followed by a ':', then add the value next to this placeholder.
        # If there is a Dear <<full name>>, make sure to write Dear followed by the name
        # Always follow data format as dd/MM/YYYY
        # - Example 1:
        #     Template --> <<LWD dd/MM/YYYY>>
        #     Your response --> 12/12/2024

        # - If the placeholder is in the form of a key (Eg. Date:), make sure to retain the placeholder and add its value next to it.
        #     Example:
        #     Template --> Date:
        #     Your response --> Date: 12/12/2024


        # If the replaced value is part of a sentence or paragraph, wrap only the value/attribute.
        # Example:
        # <<LWD dd/MM/YYYY>> → ##12/12/2024##
        # If the placeholder includes a key-value pair, wrap both the key and value together.
        # Example:
        # Date: → ##Date: 12/12/2024##


        # Do not wrap pronouns such as "his," "her," or "him," even if they reference replaced attributes.

        # Do NOT write anything after 'Dear name', or 'Name'
        # Do NOT unnessasarily add all the info in input data as key: value pair
        # """

            #     - If the placeholder is a key-value pair, wrap both the key and value together in double hashes.  
            # Example:  
            # - Template → `Date:`  
            # Response → `##Date: 12/12/2024##`  
            # Example:
            # - Template → `Employee code:`  
            # Response → `##Employee code: 165##`

        # 4.  Placeholder Replacement Scenarios :  
        # - For placeholders enclosed in `<< >>` or in the form of a blank line (`__________`): replace with the value directly.  
        #     Example:  
        #     - Template → `<<LWD dd/MM/YYYY>>`  
        #     Response → `##12/12/2024##`  
        # - For placeholders followed by a colon (`:`): retain the key and append the value.  
        #     Example:  
        #     - Template → `Date:`  
        #     Response → `##Date: 12/12/2024##`  
        # - For "Dear <<full name>>", replace with "Dear [Full Name]" without adding any additional content after the name.  

        # 5.  Text Wrapping for Further Processing :  
        # - If the replaced value is part of a sentence or paragraph, wrap the value or attribute in double hashes (`##`).  
        #     Example:  
        #     - Template → `<<LWD dd/MM/YYYY>>`  
        #     Response → `##12/12/2024##`  

        # - If the value added followed by a placeholder separated by a colon (":"), wrap both placeholder, colon, value.
        #     Example:
        #     - Template -> `Date:`
        #     Response -> `##Date: 12/12/2024##`

        # - If there is a section header followed by a colon (":") and some information, wrap the section header including the colon (":")
        #     Example:
        #     - Template -> `Confidentiality:  It is hereby agreed by the Employee that the information concerning the workings of the Company...`
        #     Response -> `##Confidentiality:##  It is hereby agreed by the Employee that the information concerning the workings of the Company`


        prompt = f"""
        You are an AI tasked with replacing placeholders in an HR document template using the provided input data.  

        Template : {context}  
        Input Data : {input_data.to_csv(index=False)}  

        ###  Instructions for Replacing Placeholders :  
        1.  General Replacement Rules :  
        - Replace placeholders with the most contextually relevant value from the input data.  
        - If no relevant value exists, leave the placeholder unchanged.  

        2.  Pronouns and Gender Formatting :  
        - Format pronouns (e.g., his/her, Mr./Mrs.) based on the gender of the name in the input data.  

        3.  Date Formatting :  
        - If a date placeholder (e.g., `<<LWD dd/MM/YYYY>>`) is present, replace it with a date in the format `dd/MM/YYYY` (e.g., `12/12/2024`).  

        4. General Placeholder Replacement with wrapping:

        - Case 1: For placeholders enclosed in `<< >>` or in the form of a blank line (__________), replace them with the value directly.
            - Include the wrapping which is enclosing the values replaced within double hashed (##)
            - In this case, do not include the placeholder name but just the value with wrapping.

            Examples:

            Template → <<LWD dd/MM/YYYY>>
            Response → ##12/12/2024##

            Template → <<DOJ dd/MM/YYYY>>
            Response → ##12/12/2024##

            Template -> We are pleased to appoint you as <<Designation>>, at our organization Affine Analytics Pvt. Ltd. with effect from <<DOJ dd/MM/YYYY>>
            Response -> We are pleased to appoint you as ##Data scientist##, at our organization Affine Analytics Pvt. Ltd. with effect from ##12/12/2024##

        - Case 2: For placeholders or section headers followed by a colon (:), retain the key (including the colon) and append the value
            - Include the wrapping which is enclosing the key and value appended within double hashed (##)
            - In this case, retain the placeholder name but just appened value after the colon (":") and everython with wrapping.

            Examples:

            Template → Date:
            Response → ##Date: 12/12/2024##

            Template → Exployee code:
            Response → ##Exployee code: 165##

            Template → Designation:
            Response → ##Designation: Data scientist##

        - Case 3: If the section header followed by a colon (:) is succeeded by content, wrap the section header (including the colon) in double hashes.
            
            Example:

            Template → Confidentiality: It is hereby agreed by the Employee that the information concerning the workings of the Company...
            Response → ##Confidentiality:## It is hereby agreed by the Employee that the information concerning the workings of the Company...

            Template - Full Time Employment: The Employee is appointed as a full time employee of the Company & shall devote his time exclusively for the business of the Company.
            Response - ##Full Time Employment:## The Employee is appointed as a full time employee of the Company & shall devote his time exclusively for the business of the Company.

        5. Do not wrap pronouns (e.g., "his," "her," "him") even if they reference replaced attributes.  

        Do NOT write anything after 'Dear name', or 'Name'
        Do NOT unnessasarily add all the info in input data as key: value pair
        Do NOT inlcude words such as template or response in your answer. That is only for example purpose.
        Answer only what is expected. Do NOT add/include any introductory or concluding phrases
        """
        try:
            # response = openai.ChatCompletion.create(
            #     model=gpt_model,
            #     messages=[{"role": "system", "content": "You are a helpful assistant."},
            #               {"role": "user", "content": prompt}],
            #     max_tokens=100,
            #     temperature=0,
            # )
            # return response["choices"][0]["message"]["content"].strip()

            completion = client.chat.completions.create(
                model=model_name,
                temperature=0.1,
                messages=[{'role': 'system', 'content': 'You are a helpful assistant.'},
                          {"role": "user", "content": prompt}])

            # messages=[{'role': 'system', 'content': 'You are a helpful assistant.'},
            #         {"role": "user", "content":
            #         [
            #     {
            #         "type": "text",
            #         "text": f"""{prompt}
            #         """
            #     },
            #     {
            #         "type": "image_url",
            #         "image_url": {
            #             "url": local_image_to_data_url(image_path='/home/affine/Projects/HR-Tools/image_example.png')
            #         }
            #     }

            #         ]

            #         }])
            output = completion.choices[0].message.content
            return output
        except Exception as e:
            print(f"Error calling GPT API: {e}")
            return placeholder  # Return the placeholder unchanged in case of error

    def get_full_name_gpt(data):
        prompt = f"""
        Read the input CSV and provide the full name:
        - If a 'Full Name' column exists, return its value.
        - If 'First Name' and 'Last Name' columns exist, combine them as 'First Name Last Name'.
        - If only 'First Name' exists, return its value.

        Input CSV:
        {data.to_csv(index=False)}

        Do NOT add any introductory or concluding phrases.
        """

        completion = client.chat.completions.create(
            model=model_name,
            temperature=0.1,
            messages=[{'role': 'system', 'content': 'You are a helpful assistant.'},
                      {"role": "user", "content": prompt}])

        output = completion.choices[0].message.content
        return output

    def get_table_value(placeholder, data):
        prompt = f"""
        Read the input CSV and provide the relevant and appropriate value for {placeholder}:
        Answer just the value.
        If you don't find relevant and approprioate value, just return NA

        Input CSV:
        {data.to_csv(index=False)}

        Do NOT add any introductory or concluding phrases.
        """

        completion = client.chat.completions.create(
            model=model_name,
            temperature=0.1,
            messages=[{'role': 'system', 'content': 'You are a helpful assistant.'},
                      {"role": "user", "content": prompt}])

        output = completion.choices[0].message.content
        return output

    def check_for_placeholders(text):
        """Check if a string contains placeholders."""
        text = text.lower().replace(' /','/').replace('/ ','/')
        #"date:", "place:", "location:",
        return any(
            placeholder in text for placeholder in ["<<", ">>", "{", "}", "__",  "his/her", "her/his", "him/her", "her/him", "mr/mrs", "mr/ms"])

    def get_signature_full_name(text):
        pattern = r"<<.*?>>"
        # Find all matches
        matches = re.findall(pattern, para.text)
        if matches:
            return matches[0]

    # Step 1: Replace placeholders in text paragraphs
    full_name = get_full_name_gpt(data=data)
    i = 1
    for para in doc.paragraphs:
        print("para:   ", i)
        print("papa text:  ", para.text)
        placeholder_present = check_for_placeholders(para.text)

        if placeholder_present:
            if '    <<' not in para.text.lower():
                gpt_output = gpt_fill_placeholder(para.text, data)
                print("gpt_output before:  \n", gpt_output)
                # gpt_output = wrap_double_quotes_with_hashes(gpt_output)
                # gpt_output = wrap_phrases_with_hashes(gpt_output)
                # print("gpt_output after:  \n", gpt_output)
                # if gpt_output.strip().lower().startswith('<<lwd'):
                #     gpt_output = gpt_output.lstrip('<<lwd').strip('>>').strip()
                para.text = gpt_output
            elif '    <<' in para.text.lower() and 'name' in para.text.lower():
                full_name_placeholder = get_signature_full_name(para.text)
                # full_name_ = '##' + full_name + '##'
                para.text ='##' + para.text.replace(full_name_placeholder, full_name) + '##'

        i += 1

    for table in doc.tables:
        # Ensure the table has exactly two columns
        if all(len(row.cells) == 2 for row in table.rows):
            for row in table.rows:
                key_cell, value_cell = row.cells
                print(key_cell.text, value_cell.text)
                # Check if the second cell (value_cell) is empty
                if not value_cell.text.strip():  # Second cell is empty
                    placeholder = key_cell.text.strip()  # Get the text from the first cell
                    # context = f"Key: {placeholder}, Value Cell is empty."
                    # Pass the placeholder and context to GPT to generate the value
                    updated_value = get_table_value(placeholder, data)
                    # Update the second cell with GPT's output
                    value_cell.text = updated_value


    for paragraph in doc.paragraphs:
        # Create a list to hold new runs
        new_runs = []
        for run in paragraph.runs:
            # Split the run text into segments based on the ## markers
            matches = re.split(r'(##.*?##)', run.text)

            for match in matches:
                if match.startswith("##") and match.endswith("##"):
                    # Extract the text inside the markers and create a bold run
                    bold_text = match.strip("#")
                    new_run = paragraph.add_run(bold_text)
                    new_run.bold = True
                elif match.startswith('"') and match.endswith('"'):
                    # Create a bold run for text enclosed in double quotes, including the quotes
                    bold_text = match
                    new_run = paragraph.add_run(bold_text)
                    new_run.bold = True
                else:
                    # Copy the formatting of the original run for non-bold text
                    new_run = paragraph.add_run(match)
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                    new_run.font.color.rgb = run.font.color.rgb
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.bold = run.bold  # Retain existing bold formatting if present
                    
                new_runs.append(new_run)

            # Clear the original text in the run
            run.text = ""

        # Add the new runs back to the paragraph
        for new_run in new_runs:
            paragraph._element.append(new_run._element)


    # Save the updated document to the specified output path
    # doc.save(output_path)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
    # print(f"Document successfully populated and saved to {"C:/Users/Lenovo/PycharmProjects/PythonProject1"}")


# import time

# st = time.time()
# if __name__ == "__main__":
#     populate_docx_with_gpt("templates/Offshore_Appointment letter.docx", "test_ouput/Book 4.xlsx",
#                            "/home/affine/Projects/HR-Tools/test_ouput/output-31.docx")

# et = time.time()

# print("TIME TAKEN:   ", round((et - st), 2))



###############
# time_taken_dt = {}
# if __name__ == "__main__":
#     templates = os.listdir('templates/')
#     for i in range(3):
#         for template in templates:
#             st = time.time()
#             output_file_name = f"iter_{i}_tem_{template}_output"
#             populate_docx_with_gpt(
#                 'templates/' +template,
#                 "test_ouput/Book 4.xlsx",
#                 f"test_ouput/{output_file_name}.docx"
#             )
#             et = time.time()
#             time_taken = round((et - st), 2)
#             time_taken_dt[output_file_name] = time_taken

#             print("TIME TAKEN:   ", time_taken)
